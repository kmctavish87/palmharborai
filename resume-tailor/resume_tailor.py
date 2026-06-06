from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
import zipfile
import os
from html import unescape
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from urllib.parse import parse_qs, urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document


APP_ROOT = Path(__file__).resolve().parent
LOCAL_GITHUB_FOLDER = Path(
    "/Users/kylemctavish/Documents/Documents - Kyle’s MacBook Air/GitHub"
)
PACKAGED_RESUME_FOLDER = APP_ROOT / "resume_templates"
DEFAULT_RESUME_SOURCE_FOLDER = (
    PACKAGED_RESUME_FOLDER if PACKAGED_RESUME_FOLDER.exists() else LOCAL_GITHUB_FOLDER
)
RESUME_SOURCE_FOLDER = Path(
    os.getenv("RESUME_SOURCE_FOLDER", str(DEFAULT_RESUME_SOURCE_FOLDER))
)
DEFAULT_OUTPUT_FOLDER = (
    LOCAL_GITHUB_FOLDER / "Generated Resumes"
    if LOCAL_GITHUB_FOLDER.exists()
    else Path(tempfile.gettempdir()) / "resume-tailor-exports"
)
OUTPUT_FOLDER = Path(os.getenv("OUTPUT_FOLDER", str(DEFAULT_OUTPUT_FOLDER)))

USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0 Safari/537.36"
)

STOP_WORDS = {
    "a",
    "about",
    "across",
    "after",
    "all",
    "also",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "can",
    "for",
    "from",
    "has",
    "have",
    "in",
    "into",
    "is",
    "it",
    "its",
    "of",
    "on",
    "or",
    "our",
    "that",
    "the",
    "their",
    "this",
    "to",
    "with",
    "will",
    "work",
    "you",
    "your",
    "company",
    "job",
    "title",
    "responsibilities",
    "qualifications",
    "requirements",
    "role",
    "apply",
    "careers",
    "department",
    "departments",
    "location",
    "locations",
    "office",
    "offices",
    "open",
    "positions",
    "see",
    "us",
}

SECTION_PATTERNS = {
    "summary": re.compile(r"(summary|profile|overview)", re.I),
    "competencies": re.compile(r"(core competencies|skills|expertise|tools)", re.I),
    "experience": re.compile(r"(professional experience|experience|work history)", re.I),
    "education": re.compile(r"(education|certifications)", re.I),
}


@dataclass
class ParsedJob:
    title: str
    company: str
    description: str
    responsibilities: list[str]
    qualifications: list[str]
    keywords: list[str]
    source_url: str = ""


@dataclass
class ResumeCandidate:
    path: Path
    text: str
    score: float
    matched_keywords: list[str]


@dataclass
class TailoredArtifacts:
    resume_path: Path
    cover_letter_path: Path | None
    selected_resume: ResumeCandidate
    job: ParsedJob
    added_keywords: list[str]
    length_estimate: str


class ResumeTailorError(Exception):
    pass


class ScrapeError(ResumeTailorError):
    pass


def extract_job_text_from_url(url: str) -> str:
    """Fetch a job URL and return readable page text."""
    cleaned_url = url.strip()
    if not cleaned_url:
        raise ScrapeError("Enter a job posting URL.")
    if not urlparse(cleaned_url).scheme:
        cleaned_url = f"https://{cleaned_url}"

    ashby_text = extract_ashby_job_text(cleaned_url)
    if ashby_text:
        return ashby_text
    greenhouse_text = extract_greenhouse_job_text(cleaned_url)
    if greenhouse_text:
        return greenhouse_text

    try:
        response = requests.get(
            cleaned_url,
            headers={"User-Agent": USER_AGENT, "Accept-Language": "en-US,en;q=0.9"},
            timeout=15,
        )
    except requests.RequestException as exc:
        raise ScrapeError(f"The job URL could not be reached: {exc}") from exc

    if response.status_code in {401, 403, 429}:
        raise ScrapeError(
            "This site blocked automated reading. Paste the job description manually."
        )
    if response.status_code >= 400:
        raise ScrapeError(f"The job URL returned HTTP {response.status_code}.")

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "nav", "footer"]):
        tag.decompose()

    structured_text = extract_structured_page_text(soup)
    main = soup.find("main") or soup.find(attrs={"role": "main"}) or soup.body or soup
    text = structured_text or " ".join(main.get_text("\n").split())
    if len(text) < 400:
        raise ScrapeError(
            "The page did not expose enough job text. Paste the description manually."
        )
    return text


def extract_ashby_job_text(url: str) -> str:
    parsed = urlparse(url)
    query = parse_qs(parsed.query)
    job_id = (query.get("ashby_jid") or [""])[0]
    if "jobs.ashbyhq.com" in parsed.netloc and parsed.path.strip("/"):
        parts = parsed.path.strip("/").split("/")
        if len(parts) >= 2:
            board_slug, job_id = parts[0], parts[1]
            return fetch_ashby_job(board_slug, job_id)
    if not job_id:
        return ""

    try:
        response = requests.get(url, headers={"User-Agent": USER_AGENT}, timeout=15)
        response.raise_for_status()
    except requests.RequestException:
        return ""

    match = re.search(r"https://jobs\.ashbyhq\.com/([^/\"']+)/embed", response.text)
    if not match:
        return ""
    return fetch_ashby_job(match.group(1), job_id)


def extract_greenhouse_job_text(url: str) -> str:
    parsed = urlparse(url)
    if "greenhouse.io" not in parsed.netloc:
        return ""

    parts = [part for part in parsed.path.strip("/").split("/") if part]
    board_slug = ""
    job_id = ""
    if parsed.netloc.startswith("boards.greenhouse.io") and len(parts) >= 2:
        board_slug, job_id = parts[0], parts[1]
    elif parsed.netloc.startswith("job-boards.greenhouse.io") and len(parts) >= 3:
        board_slug, job_id = parts[0], parts[2]

    job_id = re.sub(r"\D", "", job_id)
    if not board_slug or not job_id:
        return ""
    return fetch_greenhouse_job(board_slug, job_id)


def fetch_greenhouse_job(board_slug: str, job_id: str) -> str:
    api_url = (
        f"https://boards-api.greenhouse.io/v1/boards/{board_slug}/jobs/{job_id}"
        "?questions=false"
    )
    try:
        response = requests.get(api_url, headers={"User-Agent": USER_AGENT}, timeout=15)
        response.raise_for_status()
        payload = response.json()
    except (requests.RequestException, ValueError):
        return ""

    description = html_to_text(unescape(payload.get("content", "")))
    company = (payload.get("company_name") or board_slug_to_company(board_slug)).strip()
    location = (payload.get("location") or {}).get("name", "")
    return clean_text(
        "\n".join(
            [
                f"Company: {company}",
                f"Job Title: {(payload.get('title') or '').strip()}",
                f"Location: {location.strip()}",
                "",
                description,
            ]
        )
    )


def fetch_ashby_job(board_slug: str, job_id: str) -> str:
    api_url = f"https://api.ashbyhq.com/posting-api/job-board/{board_slug}"
    try:
        response = requests.get(api_url, headers={"User-Agent": USER_AGENT}, timeout=15)
        response.raise_for_status()
        payload = response.json()
    except (requests.RequestException, ValueError):
        return ""

    for job in payload.get("jobs", []):
        if job.get("id") != job_id:
            continue
        description = job.get("descriptionPlain") or html_to_text(
            job.get("descriptionHtml", "")
        )
        company = infer_company(description, job.get("jobUrl", "")) or board_slug_to_company(
            board_slug
        )
        return clean_text(
            "\n".join(
                [
                    f"Company: {company}",
                    f"Job Title: {job.get('title', '').strip()}",
                    f"Department: {job.get('department', '').strip()}",
                    f"Location: {job.get('location', '').strip()}",
                    "",
                    description,
                ]
            )
        )
    return ""


def extract_structured_page_text(soup: BeautifulSoup) -> str:
    chunks: list[str] = []
    for selector in [
        "[data-qa='job-title']",
        "[data-testid='job-title']",
        ".job-title",
        "h1",
        "h2",
        "[data-qa='job-description']",
        "[data-testid='job-description']",
        ".job-description",
        "[class*='jobDescription']",
    ]:
        for node in soup.select(selector):
            text = clean_text(node.get_text("\n"))
            if text and text not in chunks:
                chunks.append(text)
    return "\n\n".join(chunks)


def parse_job_description(text: str, source_url: str = "") -> ParsedJob:
    normalized = clean_text(text)
    if len(normalized) < 100:
        raise ResumeTailorError("Paste a fuller job description before continuing.")

    source_metadata = parse_source_job_metadata(source_url)
    title = (
        source_metadata.get("title", "")
        or find_label_value(normalized, ["job title", "position", "role"])
        or infer_title(normalized)
    )
    company = (
        source_metadata.get("company", "")
        or find_label_value(normalized, ["company", "organization"])
        or infer_company(normalized, source_url)
    )
    responsibilities = extract_section_items(
        normalized, ["responsibilities", "what you will do", "the role", "about the role"]
    )
    qualifications = extract_section_items(
        normalized,
        ["qualifications", "requirements", "what you bring", "skills", "experience"],
    )
    keywords = extract_keywords(normalized)

    return ParsedJob(
        title=title,
        company=company,
        description=normalized,
        responsibilities=responsibilities[:12],
        qualifications=qualifications[:12],
        keywords=keywords[:30],
        source_url=source_url,
    )


def parse_source_job_metadata(source_url: str) -> dict[str, str]:
    if not source_url.strip():
        return {}

    for extractor in (extract_greenhouse_job_text, extract_ashby_job_text):
        source_text = extractor(source_url)
        if source_text:
            return {
                "title": find_label_value(source_text, ["job title"]),
                "company": find_label_value(source_text, ["company"]),
            }
    return {}


def scan_resume_files(folder: Path = RESUME_SOURCE_FOLDER) -> list[Path]:
    if not folder.exists():
        raise ResumeTailorError(f"Resume source folder does not exist: {folder}")
    excluded_parts = {
        ".venv",
        "Generated Resumes",
        "Final PDFs",
        "Linkedin Apply",
        "LINKEDIN APPLY",
        "intentenginemarketing",
        "__pycache__",
    }
    files = []
    for path in folder.rglob("*.docx"):
        if any(part in excluded_parts for part in path.parts):
            continue
        name = path.name.lower()
        if path.name.startswith("~$") or "cover letter" in name:
            continue
        if "kyle" not in name or "mctavish" not in name:
            continue
        files.append(path)
    if not files:
        raise ResumeTailorError(f"No resume .docx files were found in {folder}")
    return sorted(files)


def extract_docx_text(path: Path) -> str:
    try:
        document = Document(path)
        chunks = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                chunks.extend(cell.text for cell in row.cells)
        return clean_text("\n".join(chunks))
    except Exception:
        return extract_docx_text_fast(path)


def select_best_resume_match(job: ParsedJob, files: Iterable[Path]) -> ResumeCandidate:
    job_terms = set(extract_keywords(job.description, limit=80))
    candidates: list[ResumeCandidate] = []
    for path in files:
        text = extract_docx_text(path)
        resume_terms = set(extract_keywords(text, limit=120))
        matched = sorted(job_terms & resume_terms)
        score = len(matched) / max(len(job_terms), 1)
        score += role_name_bonus(path.name, job.title)
        candidates.append(
            ResumeCandidate(path=path, text=text, score=score, matched_keywords=matched[:20])
        )
    return max(candidates, key=lambda item: item.score)


def build_preview(job: ParsedJob, selected: ResumeCandidate) -> dict[str, object]:
    source_keywords = set(extract_keywords(selected.text, limit=140))
    addable = [word for word in job.keywords if word in source_keywords]
    return {
        "job_title": job.title,
        "company": job.company,
        "selected_resume": str(selected.path),
        "matched_keywords": selected.matched_keywords[:12],
        "missing_keywords_added": addable[:8],
        "length_estimate": estimate_resume_length(selected.path),
    }


def export_tailored_documents(
    job: ParsedJob,
    selected: ResumeCandidate,
    include_cover_letter: bool,
    output_folder: Path = OUTPUT_FOLDER,
) -> TailoredArtifacts:
    if not output_folder.exists():
        raise ResumeTailorError(
            f"Output folder does not exist: {output_folder}. Create it and try again."
        )

    safe_company = safe_filename(job.company)
    safe_title = safe_filename(job.title)
    resume_path = output_folder / f"{safe_company}_{safe_title}_Kyle_McTavish_Resume.docx"
    cover_path = (
        output_folder
        / f"{safe_company}_{safe_title}_Kyle_McTavish_Cover_Letter.docx"
    )

    shutil.copy2(selected.path, resume_path)
    document = Document(resume_path)
    added_keywords = tailor_resume_document(document, job, selected.text)

    shortened = enforce_two_page_limit(document, job)
    if not shortened:
        raise ResumeTailorError(
            "The resume could not be kept under 2 pages without removing too much content."
        )
    document.save(resume_path)

    final_pages = count_pages_with_libreoffice(resume_path)
    if final_pages and final_pages > 2:
        raise ResumeTailorError(
            f"The exported resume is still {final_pages} pages after shortening."
        )

    generated_cover_path = None
    if include_cover_letter:
        create_cover_letter(cover_path, job, selected)
        generated_cover_path = cover_path

    return TailoredArtifacts(
        resume_path=resume_path,
        cover_letter_path=generated_cover_path,
        selected_resume=selected,
        job=job,
        added_keywords=added_keywords,
        length_estimate=estimate_resume_length(resume_path),
    )


def tailor_resume_document(document: Document, job: ParsedJob, source_text: str) -> list[str]:
    source_keywords = set(extract_keywords(source_text, limit=160))
    allowed_keywords = [keyword for keyword in job.keywords if keyword in source_keywords]
    missing_truthful = [keyword for keyword in job.keywords if keyword in source_keywords][:8]

    replace_headline(document, job, allowed_keywords)
    replace_summary(document, job, allowed_keywords)
    emphasize_skills(document, job.keywords)
    reorder_competencies(document, job.keywords)
    tailor_bullets_in_place(document, job.keywords)
    return missing_truthful


def replace_headline(document: Document, job: ParsedJob, keywords: list[str]) -> None:
    focus = build_focus_phrase(keywords)
    headline = f"{job.title} | {focus}" if focus else job.title
    for paragraph in document.paragraphs[:12]:
        text = paragraph.text.strip()
        if not text or "kyle" in text.lower() or "@" in text:
            continue
        if len(text) <= 90 and not text.endswith("."):
            set_paragraph_text(paragraph, headline)
            return


def replace_summary(document: Document, job: ParsedJob, keywords: list[str]) -> None:
    summary_index = find_section_index(document.paragraphs, "summary")
    target = (
        next_content_paragraph(document.paragraphs, summary_index + 1)
        if summary_index is not None
        else first_profile_paragraph(document.paragraphs)
    )
    if target is None:
        return

    supported = prioritized_supported_terms(keywords)
    keyword_phrase = ", ".join(supported[:7])
    replacement = (
        "Demand generation and revenue marketing leader with 15+ years of "
        "experience driving pipeline growth, full-funnel campaign performance, "
        "ABM, marketing automation, sales alignment, attribution, and GTM "
        "execution across SaaS, consulting, telecommunications, and enterprise "
        f"environments. Tailored for {job.company}'s {job.title} role"
    )
    if keyword_phrase:
        replacement += f", with strongest alignment around {keyword_phrase}."
    else:
        replacement += "."
    set_paragraph_text(target, replacement)


def emphasize_skills(document: Document, job_keywords: list[str]) -> None:
    replacements = {
        "Pipeline strategy & revenue growth": "Pipeline strategy, revenue growth & sales alignment",
        "Full-funnel demand generation & conversion optimization": "Full-funnel demand generation, campaign performance & conversion optimization",
        "Account-Based Marketing (ABM) & intent-driven targeting": "Account-Based Marketing (ABM), intent-driven targeting & enterprise segmentation",
        "Cross-portfolio GTM planning & campaign orchestration": "GTM planning, campaign orchestration & cross-functional execution",
        "Attribution modeling, analytics & performance optimization": "Attribution modeling, funnel analytics & performance optimization",
    }
    for paragraph in document.paragraphs[:25]:
        text = paragraph.text.strip()
        if text in replacements and keyword_relevance(text, job_keywords) > 0:
            set_paragraph_text(paragraph, replacements[text])


def tailor_bullets_in_place(document: Document, job_keywords: list[str]) -> None:
    for paragraph in document.paragraphs:
        if not is_bullet_paragraph(paragraph):
            continue
        text = paragraph.text.strip()
        replacement = rewrite_supported_bullet(text, job_keywords)
        if replacement != text:
            set_paragraph_text(paragraph, replacement)


def rewrite_supported_bullet(text: str, job_keywords: list[str]) -> str:
    normalized = text.strip()
    lower = normalized.lower()
    if "managed entirely wipfli" in lower:
        return (
            "Managed Wipfli's AI services and solutions marketing focus, supporting "
            "demand generation and pipeline growth initiatives."
        )
    if "designed and executed full-funnel demand generation programs" in lower:
        return (
            "Designed and executed full-funnel demand generation programs supporting "
            "GTM strategy, pipeline acceleration, and campaign performance."
        )
    if "built and optimized abm" in lower:
        return (
            "Built and optimized ABM and intent-driven campaigns focused on high-value "
            "accounts, segmentation, and pipeline creation."
        )
    if "partnered with sales leadership" in lower:
        return (
            "Partnered with sales leadership to improve MQL-to-SQL conversion, funnel "
            "velocity, and sales-aligned demand generation execution."
        )
    if "enhanced attribution" in lower:
        return (
            "Enhanced attribution and measurement frameworks to clarify marketing "
            "influence, campaign performance, and funnel progression."
        )
    if "revamped entire digital marketing strategy" in lower:
        return (
            "Revamped digital marketing strategy to improve channel efficiency, lower "
            "CPC, and strengthen demand generation performance."
        )
    if "exceeded sql targets" in lower:
        return (
            "Exceeded SQL targets by improving campaign execution and demand generation "
            "program performance in partnership with sales."
        )
    if "generated over $17.6m in pipeline" in lower:
        return (
            "Generated over $17.6M in pipeline by leading demand generation, digital, "
            "automation, and event marketing initiatives across North America."
        )
    if "cultivated and maintained business relations" in lower:
        return (
            "Managed key vendor relationships across Salesforce, Marketo, and Demandbase "
            "to support revenue marketing operations and campaign execution."
        )
    if "introduced and rolled out proprietary" in lower:
        return (
            "Introduced marketing effectiveness measurement tools to improve market "
            "planning, demand generation, and performance visibility."
        )
    if keyword_relevance(normalized, job_keywords) >= 2 and len(normalized) > 155:
        return trim_sentence(normalized, 150)
    return normalized


def build_focus_phrase(keywords: list[str]) -> str:
    supported = prioritized_supported_terms(keywords)
    if not supported:
        return "Pipeline Growth, ABM & Revenue Marketing"
    preferred = []
    for term in [
        "pipeline",
        "demand",
        "generation",
        "sales",
        "campaign",
        "abm",
        "attribution",
        "marketing",
    ]:
        if term in supported and term not in preferred:
            preferred.append(term)
    if "demand" in preferred and "generation" in preferred:
        preferred = ["demand generation" if term == "demand" else term for term in preferred]
        preferred = [term for term in preferred if term != "generation"]
    labels = {
        "pipeline": "Pipeline Growth",
        "demand generation": "Demand Generation",
        "sales": "Sales Alignment",
        "campaign": "Campaign Strategy",
        "abm": "ABM",
        "attribution": "Attribution",
        "marketing": "Revenue Marketing",
    }
    return ", ".join(labels.get(term, term.title()) for term in preferred[:3])


def prioritized_supported_terms(keywords: list[str]) -> list[str]:
    preferred_order = [
        "pipeline",
        "demand",
        "generation",
        "campaign",
        "sales",
        "abm",
        "attribution",
        "automation",
        "marketing",
        "gtm",
        "revenue",
    ]
    result = [term for term in preferred_order if term in keywords]
    result.extend(term for term in keywords if term not in result)
    return result
    set_paragraph_text(target, replacement)


def reorder_competencies(document: Document, job_keywords: list[str]) -> None:
    section_index = find_section_index(document.paragraphs, "competencies")
    if section_index is None:
        return

    for paragraph in document.paragraphs[section_index + 1 : section_index + 5]:
        text = paragraph.text.strip()
        if not text or is_section_heading(text):
            break
        separators = ";" if ";" in text else "|"
        if separators not in text and "," in text:
            separators = ","
        items = [item.strip() for item in re.split(r"[;|,]", text) if item.strip()]
        if len(items) < 4:
            continue
        ranked = sorted(
            items,
            key=lambda item: keyword_relevance(item, job_keywords),
            reverse=True,
        )
        set_paragraph_text(paragraph, f" {separators} ".join(ranked))
        return


def reprioritize_bullets(document: Document, job_keywords: list[str]) -> None:
    bullets = [
        paragraph
        for paragraph in document.paragraphs
        if is_bullet_paragraph(paragraph)
    ]
    if len(bullets) < 4:
        return

    ranked_text = sorted(
        [paragraph.text.strip() for paragraph in bullets],
        key=lambda text: keyword_relevance(text, job_keywords),
        reverse=True,
    )
    for paragraph, text in zip(bullets, ranked_text):
        set_paragraph_text(paragraph, lightly_rewrite_bullet(text, job_keywords))


def lightly_rewrite_bullet(text: str, job_keywords: list[str]) -> str:
    text = re.sub(r"^[•\-–]\s*", "", text).strip()
    if not text:
        return text
    strongest = [word for word in job_keywords if re.search(rf"\b{re.escape(word)}\b", text, re.I)]
    if strongest and not text.lower().startswith(strongest[0].lower()):
        return text
    return text


def enforce_two_page_limit(document: Document, job: ParsedJob) -> bool:
    """Shorten conservatively using a character/page proxy before saving."""
    max_chars = 9300
    if estimate_document_chars(document) <= max_chars:
        return True

    bullets = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip()
        and (
            paragraph.text.strip().startswith(("•", "-", "–"))
            or "List" in (paragraph.style.name if paragraph.style else "")
        )
    ]
    for paragraph in sorted(
        bullets, key=lambda p: keyword_relevance(p.text, job.keywords)
    ):
        if estimate_document_chars(document) <= max_chars:
            return True
        shorten_or_clear_bullet(paragraph)

    for paragraph in reversed(document.paragraphs):
        if estimate_document_chars(document) <= max_chars:
            return True
        text = paragraph.text.strip()
        if text and not is_section_heading(text) and "kyle" not in text.lower():
            set_paragraph_text(paragraph, trim_sentence(text, 140))

    return estimate_document_chars(document) <= max_chars


def create_cover_letter(path: Path, job: ParsedJob, selected: ResumeCandidate) -> None:
    document = Document()
    document.add_paragraph("Kyle McTavish")
    document.add_paragraph("")
    document.add_paragraph(f"Dear {job.company} Hiring Team,")
    document.add_paragraph("")
    matched = ", ".join(selected.matched_keywords[:6])
    document.add_paragraph(
        f"I am excited to apply for the {job.title} role at {job.company}. "
        f"My background aligns closely with the needs outlined in the posting, "
        f"especially around {matched or 'growth, marketing operations, and revenue execution'}."
    )
    document.add_paragraph(
        "Across my resume examples, I have emphasized measurable growth programs, "
        "cross-functional go-to-market execution, lifecycle strategy, campaign "
        "operations, and practical use of marketing technology. I would welcome "
        "the opportunity to bring that experience to your team."
    )
    document.add_paragraph(
        "Thank you for your consideration. I would be glad to discuss how my "
        "experience can support your goals."
    )
    document.add_paragraph("")
    document.add_paragraph("Sincerely,")
    document.add_paragraph("Kyle McTavish")
    document.save(path)


def clean_text(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", re.sub(r"[ \t]+", " ", text)).strip()


def find_label_value(text: str, labels: list[str]) -> str:
    for label in labels:
        match = re.search(
            rf"(?:^|\n)\s*{re.escape(label)}\s*[:\-]\s*([^\n|•]+)",
            text,
            re.I,
        )
        if match:
            return tidy_title(match.group(1)[:80])
    return ""


def infer_title(text: str) -> str:
    lines = [line.strip() for line in re.split(r"[\n|]", text) if line.strip()]
    derived = infer_unlabeled_role_title(text)
    if derived:
        return derived
    role_words = re.compile(
        r"\b(manager|director|lead|specialist|analyst|strategist|consultant|"
        r"engineer|operator|owner|head|vp|president|marketer)\b",
        re.I,
    )
    rejected_copy = re.compile(
        r"(turn ai|not just|experiments|systems i build|outcomes|clients would say|"
        r"let.?s build|qualified leads|booked jobs|client retention|time and cost|"
        r"automations and internal tools|how you.?ll be measured|compensation|benefits)",
        re.I,
    )
    for line in lines[:25]:
        if 5 <= len(line) <= 90 and role_words.search(line) and not rejected_copy.search(line):
            return tidy_title(line)
    return "Target Role"


def infer_unlabeled_role_title(text: str) -> str:
    lower = text.lower()
    if "ai-native marketer" in lower:
        if "portfolio of accounts" in lower or "account outcomes" in lower:
            return "AI-Native Account Marketer"
        return "AI-Native Marketer"
    if "performance marketing" in lower and "lead generation" in lower:
        return "Performance Marketing Lead"
    if "growth marketing" in lower and "lead generation" in lower:
        return "Growth Marketing Lead"
    return ""


def infer_company(text: str, url: str) -> str:
    first_line_company = re.match(
        r"^([A-Z][A-Za-z0-9&. ]{2,40}?)\s+"
        r"(?:empowers|is|provides|builds|helps|creates|delivers)\b",
        text.strip(),
    )
    if first_line_company:
        return tidy_title(first_line_company.group(1))

    match = re.search(
        r"\bat\s+([A-Za-z][A-Za-z0-9&. ]{2,50}?)(?:,|\s+you\b|\s+we\b|\.|\n)",
        text,
    )
    if match:
        return tidy_title(match.group(1))
    parsed = urlparse(url if urlparse(url).scheme else f"https://{url}")
    domain = parsed.netloc.replace("www.", "")
    parts = [part for part in parsed.path.strip("/").split("/") if part]
    if "greenhouse.io" in domain and parts:
        return board_slug_to_company(parts[0])
    if domain and domain not in {"jobs.ashbyhq.com", "api.ashbyhq.com"}:
        return tidy_title(domain_to_company(domain))
    return ""


def board_slug_to_company(slug: str) -> str:
    spaced = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", slug).strip()
    return tidy_title(spaced) or "Company"


def domain_to_company(domain: str) -> str:
    name = domain.split(".")[0]
    known = {"palmharborai": "Palm Harbor AI"}
    if name.lower() in known:
        return known[name.lower()]
    return re.sub(r"(?<=[a-z])(?=[A-Z])", " ", name).title()


def extract_section_items(text: str, headings: list[str]) -> list[str]:
    pattern = "|".join(re.escape(heading) for heading in headings)
    match = re.search(rf"({pattern})\s*[:\n](.+?)(\n[A-Z][A-Za-z ]{{3,40}}\s*[:\n]|$)", text, re.I | re.S)
    source = match.group(2) if match else text
    items = re.split(r"(?:\n|•|\s-\s|\s\*\s)", source)
    return [clean_text(item) for item in items if 35 <= len(clean_text(item)) <= 300]


def extract_keywords(text: str, limit: int = 40) -> list[str]:
    lower = text.lower()
    phrases = re.findall(r"\b[a-z][a-z0-9+#.&-]*(?:\s+[a-z][a-z0-9+#.&-]*){1,3}\b", lower)
    words = re.findall(r"\b[a-z][a-z0-9+#.&-]{2,}\b", lower)
    counts: Counter[str] = Counter()
    for phrase in phrases:
        tokens = [token for token in phrase.split() if token not in STOP_WORDS]
        if len(tokens) >= 2:
            counts[" ".join(tokens)] += 2
    for word in words:
        if word not in STOP_WORDS and not word.isdigit():
            counts[word] += 1
    return [term for term, _ in counts.most_common(limit)]


def role_name_bonus(filename: str, title: str) -> float:
    file_terms = set(extract_keywords(filename.replace("_", " "), limit=20))
    title_terms = set(extract_keywords(title, limit=20))
    return 0.12 * len(file_terms & title_terms)


def keyword_relevance(text: str, keywords: list[str]) -> int:
    text_lower = text.lower()
    return sum(1 for keyword in keywords if keyword.lower() in text_lower)


def find_section_index(paragraphs, key: str) -> int | None:
    pattern = SECTION_PATTERNS[key]
    for index, paragraph in enumerate(paragraphs):
        if pattern.fullmatch(paragraph.text.strip()) or pattern.search(paragraph.text.strip()):
            return index
    return None


def next_content_paragraph(paragraphs, start: int):
    for paragraph in paragraphs[start:]:
        text = paragraph.text.strip()
        if text and not is_section_heading(text):
            return paragraph
    return None


def first_profile_paragraph(paragraphs):
    for paragraph in paragraphs[:8]:
        text = paragraph.text.strip()
        if len(text) > 140 and not is_section_heading(text):
            return paragraph
    return None


def is_bullet_paragraph(paragraph) -> bool:
    return paragraph.text.strip().startswith(("•", "-", "–")) or "List" in (
        paragraph.style.name if paragraph.style else ""
    )


def is_section_heading(text: str) -> bool:
    if len(text) > 55:
        return False
    if any(pattern.search(text) for pattern in SECTION_PATTERNS.values()):
        return True
    return text.isupper() and len(text.split()) <= 5


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def estimate_document_chars(document: Document) -> int:
    total = sum(len(paragraph.text) for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            total += sum(len(cell.text) for cell in row.cells)
    return total


def estimate_resume_length(path: Path) -> str:
    pages = count_pages_with_libreoffice(path)
    if pages:
        return f"{pages} page{'s' if pages != 1 else ''} by LibreOffice export"
    chars = len(extract_docx_text(path))
    estimated_pages = max(1, round(chars / 4100, 1))
    return f"about {estimated_pages} pages by text-density estimate"


def count_pages_with_libreoffice(path: Path) -> int | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    with tempfile.TemporaryDirectory() as tmpdir:
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                tmpdir,
                str(path),
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=45,
            check=False,
        )
        if result.returncode != 0:
            return None
        pdf = Path(tmpdir) / f"{path.stem}.pdf"
        if not pdf.exists():
            return None
        data = pdf.read_bytes()
        return max(data.count(b"/Type /Page"), data.count(b"/Page")) or None


def shorten_or_clear_bullet(paragraph) -> None:
    text = paragraph.text.strip()
    if len(text) > 150:
        set_paragraph_text(paragraph, trim_sentence(text, 120))
    else:
        set_paragraph_text(paragraph, "")


def trim_sentence(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    trimmed = text[:max_chars].rsplit(" ", 1)[0].rstrip(" ,;")
    return f"{trimmed}."


def safe_filename(value: str) -> str:
    value = re.sub(r"[^\w\s.-]", "", value).strip()
    value = re.sub(r"\s+", "_", value)
    return value[:80] or "Untitled"


def tidy_title(value: str) -> str:
    value = clean_text(value)
    value = re.sub(r"\s+[-|].*$", "", value)
    return value.strip(" .,-_|")[:80] or "Untitled"


def extract_docx_text_fast(path: Path) -> str:
    try:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    except Exception:
        return ""
    text = re.sub(r"<[^>]+>", " ", xml)
    return clean_text(text)


def html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    return clean_text(soup.get_text("\n"))
